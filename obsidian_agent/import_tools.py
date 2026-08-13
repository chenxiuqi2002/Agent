import hashlib
import io
import os
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import NotRequired, TypedDict

from claude_agent_sdk import tool

from .config import VAULT_PATH, ATTACHMENTS_PATH, IMPORT_EXTENSIONS
from .audit import audit_logger
from .tools import _resolve_path, _build_frontmatter


def _parse_pdf(file_path: Path) -> tuple[str, list[dict]]:
    text = ""
    images = []
    try:
        import fitz
        doc = fitz.open(str(file_path))
        for page_num in range(len(doc)):
            page = doc[page_num]
            text += page.get_text() + "\n\n"
            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                try:
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image.get("ext", "png")
                    image_hash = hashlib.md5(image_bytes).hexdigest()[:12]
                    image_filename = f"{file_path.stem}_p{page_num+1}_img{img_index+1}_{image_hash}.{image_ext}"
                    images.append({
                        "filename": image_filename,
                        "data": image_bytes,
                        "ext": image_ext,
                        "page": page_num + 1,
                    })
                except Exception:
                    continue
        doc.close()
    except ImportError:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(file_path))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        except ImportError:
            text = f"[PDF parsing unavailable - install PyMuPDF or PyPDF2]\nSource: {file_path.name}"
    return text.strip(), images


def _parse_docx(file_path: Path) -> tuple[str, list[dict]]:
    text = ""
    images = []
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc = Document(str(file_path))

        for para in doc.paragraphs:
            if para.text.strip():
                style = para.style.name if para.style else ""
                if "Heading 1" in style:
                    text += f"# {para.text}\n\n"
                elif "Heading 2" in style:
                    text += f"## {para.text}\n\n"
                elif "Heading 3" in style:
                    text += f"### {para.text}\n\n"
                else:
                    text += f"{para.text}\n\n"

        for table in doc.tables:
            text += "\n| "
            for cell in table.rows[0].cells:
                text += f"{cell.text} | "
            text += "\n| "
            for _ in table.rows[0].cells:
                text += "--- | "
            text += "\n"
            for row in table.rows[1:]:
                text += "| "
                for cell in row.cells:
                    text += f"{cell.text} | "
                text += "\n"
            text += "\n"

        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    image_ext = image_part.content_type.split("/")[-1]
                    if image_ext == "jpeg":
                        image_ext = "jpg"
                    image_hash = hashlib.md5(image_bytes).hexdigest()[:12]
                    image_filename = f"{file_path.stem}_img_{image_hash}.{image_ext}"
                    images.append({
                        "filename": image_filename,
                        "data": image_bytes,
                        "ext": image_ext,
                    })
                except Exception:
                    continue

    except ImportError:
        text = f"[DOCX parsing unavailable - install python-docx]\nSource: {file_path.name}"

    return text.strip(), images


def _save_images(images: list[dict], note_stem: str) -> list[str]:
    ATTACHMENTS_PATH.mkdir(parents=True, exist_ok=True)
    saved = []
    for img in images:
        filename = img["filename"]
        dest = ATTACHMENTS_PATH / filename
        dest.write_bytes(img["data"])
        obsidian_ref = f"![[{filename}]]"
        saved.append(obsidian_ref)
    return saved


def _import_single_file(file_path: Path, output_folder: str, extract_images: bool) -> dict:
    ext = file_path.suffix.lower()
    if ext not in IMPORT_EXTENSIONS:
        return {"error": f"Unsupported format: {ext}"}

    if ext == ".pdf":
        content, images = _parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        content, images = _parse_docx(file_path)
    elif ext == ".txt":
        content = file_path.read_text(encoding="utf-8", errors="replace")
        images = []
    elif ext == ".rtf":
        try:
            from striprtf.striprtf import rtf_to_text
            rtf_content = file_path.read_text(encoding="utf-8", errors="replace")
            content = rtf_to_text(rtf_content)
        except ImportError:
            content = file_path.read_text(encoding="utf-8", errors="replace")
        images = []
    else:
        return {"error": f"Unsupported format: {ext}"}

    image_refs = []
    if extract_images and images:
        image_refs = _save_images(images, file_path.stem)

    output_name = file_path.stem + ".md"
    output_rel = os.path.join(output_folder, output_name) if output_folder else output_name
    output_path = _resolve_path(output_rel)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    metadata = {
        "title": file_path.stem,
        "tags": ["导入", ext.lstrip(".")],
        "source": str(file_path),
        "imported": datetime.now().isoformat(),
        "created": datetime.now().isoformat(),
        "updated": datetime.now().isoformat(),
    }

    full_content = _build_frontmatter(metadata) + content

    if image_refs:
        full_content += "\n\n## 附件\n\n"
        for ref in image_refs:
            full_content += f"{ref}\n"

    output_path.write_text(full_content, encoding="utf-8")

    return {
        "source": str(file_path),
        "output": output_rel,
        "images_extracted": len(image_refs),
        "content_length": len(content),
    }


class ImportFileInput(TypedDict):
    file_path: str
    output_folder: NotRequired[str]
    extract_images: NotRequired[bool]


@tool("import_file", "Import a file (PDF, DOCX, TXT, RTF) into the vault as a Markdown note. Extracts text, images, and tables.", ImportFileInput)
async def import_file(args):
    import time
    start = time.time()

    file_path_str = args["file_path"]
    output_folder = args.get("output_folder", "导入")
    extract_images = args.get("extract_images", True)

    file_path = Path(file_path_str)
    if not file_path.exists():
        result = f"File not found: {file_path_str}"
        audit_logger.log_operation("import_file", "import_file", args, result, duration_ms=(time.time() - start) * 1000, status="error")
        return {"content": [{"type": "text", "text": result}]}

    if file_path.suffix.lower() not in IMPORT_EXTENSIONS:
        result = f"Unsupported file format: {file_path.suffix}. Supported: {', '.join(IMPORT_EXTENSIONS)}"
        audit_logger.log_operation("import_file", "import_file", args, result, duration_ms=(time.time() - start) * 1000, status="error")
        return {"content": [{"type": "text", "text": result}]}

    try:
        result_data = _import_single_file(file_path, output_folder, extract_images)
        result_msg = f"File imported successfully!\nSource: {result_data['source']}\nOutput: {result_data['output']}\nImages extracted: {result_data.get('images_extracted', 0)}\nContent length: {result_data.get('content_length', 0)} chars"
        audit_logger.log_operation("import_file", "import_file", args, result_msg, duration_ms=(time.time() - start) * 1000)
        return {"content": [{"type": "text", "text": result_msg}]}
    except Exception as e:
        result = f"Import failed: {e}"
        audit_logger.log_operation("import_file", "import_file", args, result, duration_ms=(time.time() - start) * 1000, status="error", error=str(e))
        return {"content": [{"type": "text", "text": result}]}


class ImportFolderInput(TypedDict):
    source_folder: str
    output_folder: NotRequired[str]
    extract_images: NotRequired[bool]
    recursive: NotRequired[bool]


@tool("import_folder", "Import all supported files from a folder into the vault as Markdown notes.", ImportFolderInput)
async def import_folder(args):
    import time
    start = time.time()

    source_folder = args["source_folder"]
    output_folder = args.get("output_folder", "导入")
    extract_images = args.get("extract_images", True)
    recursive = args.get("recursive", True)

    source_path = Path(source_folder)
    if not source_path.exists():
        result = f"Source folder not found: {source_folder}"
        audit_logger.log_operation("import_folder", "import_folder", args, result, duration_ms=(time.time() - start) * 1000, status="error")
        return {"content": [{"type": "text", "text": result}]}

    pattern = "**/*" if recursive else "*"
    results = []
    for f in sorted(source_path.glob(pattern)):
        if f.is_file() and f.suffix.lower() in IMPORT_EXTENSIONS:
            try:
                r = _import_single_file(f, output_folder, extract_images)
                results.append(r)
            except Exception as e:
                results.append({"source": str(f), "error": str(e)})

    success = sum(1 for r in results if "error" not in r)
    failed = sum(1 for r in results if "error" in r)
    total_images = sum(r.get("images_extracted", 0) for r in results if "images_extracted" in r)

    result_msg = f"Import complete: {success} succeeded, {failed} failed, {total_images} images extracted"
    audit_logger.log_operation("import_folder", "import_folder", args, result_msg, duration_ms=(time.time() - start) * 1000)
    return {"content": [{"type": "text", "text": result_msg}]}


class ExtractImagesInput(TypedDict):
    file_path: str
    output_folder: NotRequired[str]


@tool("extract_images", "Extract images from a PDF or DOCX file and save them to the vault attachments folder.", ExtractImagesInput)
async def extract_images(args):
    import time
    start = time.time()

    file_path_str = args["file_path"]
    file_path = Path(file_path_str)

    if not file_path.exists():
        return {"content": [{"type": "text", "text": f"File not found: {file_path_str}"}]}

    ext = file_path.suffix.lower()
    if ext == ".pdf":
        _, images = _parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        _, images = _parse_docx(file_path)
    else:
        return {"content": [{"type": "text", "text": f"Image extraction not supported for: {ext}"}]}

    if not images:
        return {"content": [{"type": "text", "text": "No images found in the file."}]}

    refs = _save_images(images, file_path.stem)

    result_msg = f"Extracted {len(images)} images:\n" + "\n".join(f"  {ref}" for ref in refs)
    audit_logger.log_operation("extract_images", "extract_images", args, result_msg, duration_ms=(time.time() - start) * 1000)
    return {"content": [{"type": "text", "text": result_msg}]}


IMPORT_TOOLS = [import_file, import_folder, extract_images]
IMPORT_TOOL_NAMES = [f"mcp__obsidian__{t.name}" for t in IMPORT_TOOLS]
